import os
from reportlab.lib.pagesizes import letter
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib import colors
from reportlab.lib.enums import TA_LEFT, TA_CENTER, TA_RIGHT
from typing import Dict
from docx import Document
from docx.shared import Pt, RGBColor, Inches
from docx.enum.text import WD_ALIGN_PARAGRAPH


class DocumentGenerator:
    """
    Generate professional documents:
    1. Tailored CV (PDF and DOCX)
    2. Cover letter (PDF and DOCX)
    
    ATS-friendly formatting
    """
    
    def __init__(self):
        self.output_dir = "temp"
        os.makedirs(self.output_dir, exist_ok=True)
    
    def generate_cv(self, ai_result: Dict) -> str:
        """Generate tailored CV in PDF format"""
        
        pdf_path = os.path.join(self.output_dir, "cv.pdf")
        
        # Create PDF document
        doc = SimpleDocTemplate(pdf_path, pagesize=letter,
                              leftMargin=0.75*inch, rightMargin=0.75*inch,
                              topMargin=0.75*inch, bottomMargin=0.75*inch)
        
        # Container for PDF elements
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        title_style = ParagraphStyle(
            'CustomTitle',
            parent=styles['Heading1'],
            fontSize=24,
            textColor=colors.HexColor('#0a66c2'),
            spaceAfter=6,
            alignment=TA_CENTER
        )
        
        heading_style = ParagraphStyle(
            'CustomHeading',
            parent=styles['Heading2'],
            fontSize=14,
            textColor=colors.HexColor('#0a66c2'),
            spaceAfter=6,
            spaceBefore=12,
            borderWidth=0,
            borderPadding=0,
            borderColor=colors.HexColor('#0a66c2'),
            borderRadius=None,
        )
        
        # Name and headline
        profile_data = ai_result["profile_data"]
        story.append(Paragraph(profile_data.get("name", ""), title_style))
        story.append(Paragraph(profile_data.get("headline", ""), styles['Normal']))
        story.append(Spacer(1, 0.2*inch))
        
        # Professional Summary
        story.append(Paragraph("PROFESSIONAL SUMMARY", heading_style))
        summary = f"Experienced professional with expertise in {', '.join(ai_result.get('tailored_skills', [])[:5])}."
        story.append(Paragraph(summary, styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        # Skills
        story.append(Paragraph("SKILLS", heading_style))
        skills_text = " • ".join(ai_result.get('tailored_skills', [])[:15])
        story.append(Paragraph(skills_text, styles['Normal']))
        story.append(Spacer(1, 0.1*inch))
        
        # Experience
        story.append(Paragraph("PROFESSIONAL EXPERIENCE", heading_style))
        for exp in ai_result.get('tailored_experience', []):
            # Company and position
            company_text = f"<b>{exp.get('position', '')}</b> | {exp.get('company', '')}"
            story.append(Paragraph(company_text, styles['Normal']))
            
            # Duration
            story.append(Paragraph(exp.get('duration', ''), styles['Normal']))
            
            # Description
            desc = exp.get('description', '')
            if desc:
                story.append(Paragraph(desc, styles['Normal']))
            
            story.append(Spacer(1, 0.1*inch))
        
        # Education
        if profile_data.get('education'):
            story.append(Paragraph("EDUCATION", heading_style))
            for edu in profile_data['education']:
                edu_text = f"<b>{edu.get('degree', '')}</b> | {edu.get('institution', '')}"
                story.append(Paragraph(edu_text, styles['Normal']))
                story.append(Paragraph(edu.get('duration', ''), styles['Normal']))
                story.append(Spacer(1, 0.1*inch))
        
        # Certifications
        if profile_data.get('certifications'):
            story.append(Paragraph("CERTIFICATIONS", heading_style))
            for cert in profile_data['certifications']:
                cert_text = f"• <b>{cert.get('name', '')}</b> - {cert.get('issuer', '')}"
                story.append(Paragraph(cert_text, styles['Normal']))
            story.append(Spacer(1, 0.1*inch))
        
        # Build PDF
        doc.build(story)
        
        # Also generate DOCX version
        self._generate_cv_docx(ai_result)
        
        return pdf_path
    
    def _generate_cv_docx(self, ai_result: Dict) -> str:
        """Generate CV in DOCX format for easier editing"""
        
        docx_path = os.path.join(self.output_dir, "cv.docx")
        doc = Document()
        
        # Styles
        profile_data = ai_result["profile_data"]
        
        # Name
        name = doc.add_heading(profile_data.get("name", ""), level=1)
        name.alignment = WD_ALIGN_PARAGRAPH.CENTER
        name.runs[0].font.color.rgb = RGBColor(10, 102, 194)
        
        # Headline
        headline = doc.add_paragraph(profile_data.get("headline", ""))
        headline.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Professional Summary
        doc.add_heading('PROFESSIONAL SUMMARY', level=2)
        summary = f"Experienced professional with expertise in {', '.join(ai_result.get('tailored_skills', [])[:5])}."
        doc.add_paragraph(summary)
        
        # Skills
        doc.add_heading('SKILLS', level=2)
        skills_text = " • ".join(ai_result.get('tailored_skills', [])[:15])
        doc.add_paragraph(skills_text)
        
        # Experience
        doc.add_heading('PROFESSIONAL EXPERIENCE', level=2)
        for exp in ai_result.get('tailored_experience', []):
            p = doc.add_paragraph()
            p.add_run(f"{exp.get('position', '')}").bold = True
            p.add_run(f" | {exp.get('company', '')}")
            
            doc.add_paragraph(exp.get('duration', ''))
            doc.add_paragraph(exp.get('description', ''))
        
        # Education
        if profile_data.get('education'):
            doc.add_heading('EDUCATION', level=2)
            for edu in profile_data['education']:
                p = doc.add_paragraph()
                p.add_run(f"{edu.get('degree', '')}").bold = True
                p.add_run(f" | {edu.get('institution', '')}")
                doc.add_paragraph(edu.get('duration', ''))
        
        # Certifications
        if profile_data.get('certifications'):
            doc.add_heading('CERTIFICATIONS', level=2)
            for cert in profile_data['certifications']:
                doc.add_paragraph(f"• {cert.get('name', '')} - {cert.get('issuer', '')}")
        
        doc.save(docx_path)
        return docx_path
    
    def generate_cover_letter(self, ai_result: Dict) -> str:
        """Generate cover letter in PDF format"""
        
        pdf_path = os.path.join(self.output_dir, "cover_letter.pdf")
        
        # Create PDF document
        doc = SimpleDocTemplate(pdf_path, pagesize=letter,
                              leftMargin=1*inch, rightMargin=1*inch,
                              topMargin=1*inch, bottomMargin=1*inch)
        
        story = []
        styles = getSampleStyleSheet()
        
        # Custom styles
        header_style = ParagraphStyle(
            'Header',
            parent=styles['Normal'],
            fontSize=12,
            textColor=colors.HexColor('#0a66c2'),
            spaceAfter=20
        )
        
        # Sender info
        profile_data = ai_result["profile_data"]
        story.append(Paragraph(profile_data.get("name", ""), header_style))
        story.append(Spacer(1, 0.5*inch))
        
        # Date
        from datetime import datetime
        story.append(Paragraph(datetime.now().strftime("%B %d, %Y"), styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Recipient
        story.append(Paragraph("Hiring Manager", styles['Normal']))
        story.append(Spacer(1, 0.3*inch))
        
        # Cover letter content
        cover_letter_text = ai_result.get("cover_letter", "")
        for paragraph in cover_letter_text.split('\n\n'):
            if paragraph.strip():
                story.append(Paragraph(paragraph.strip(), styles['Normal']))
                story.append(Spacer(1, 0.15*inch))
        
        # Build PDF
        doc.build(story)
        
        # Also generate DOCX version
        self._generate_cover_letter_docx(ai_result)
        
        return pdf_path
    
    def _generate_cover_letter_docx(self, ai_result: Dict) -> str:
        """Generate cover letter in DOCX format"""
        
        docx_path = os.path.join(self.output_dir, "cover_letter.docx")
        doc = Document()
        
        profile_data = ai_result["profile_data"]
        
        # Sender info
        sender = doc.add_paragraph(profile_data.get("name", ""))
        sender.runs[0].font.color.rgb = RGBColor(10, 102, 194)
        sender.runs[0].font.size = Pt(12)
        
        doc.add_paragraph()
        
        # Date
        from datetime import datetime
        doc.add_paragraph(datetime.now().strftime("%B %d, %Y"))
        doc.add_paragraph()
        
        # Recipient
        doc.add_paragraph("Hiring Manager")
        doc.add_paragraph()
        
        # Cover letter content
        cover_letter_text = ai_result.get("cover_letter", "")
        for paragraph in cover_letter_text.split('\n\n'):
            if paragraph.strip():
                doc.add_paragraph(paragraph.strip())
        
        doc.save(docx_path)
        return docx_path
